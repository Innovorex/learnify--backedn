"""
File Processing Service
Handles extraction of text from PDF, DOCX, and TXT files
"""

import os
import re
import uuid
from pathlib import Path
from typing import Optional, Tuple, Dict
from fastapi import UploadFile, HTTPException


class FileProcessorService:
    """
    Service for processing uploaded files (PDF, DOCX, TXT)
    Handles text extraction, cleaning, and validation
    """

    UPLOAD_DIR = Path("uploads/teaching_materials")
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB

    ALLOWED_EXTENSIONS = {
        "pdf": ["application/pdf"],
        "docx": [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ],
        "txt": ["text/plain"]
    }

    def __init__(self):
        self.UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

    async def save_uploaded_file(
        self,
        file: UploadFile,
        teacher_id: int
    ) -> Tuple[str, str, int]:
        """
        Save uploaded file to disk
        Returns: (file_path, filename, file_size)
        """
        # Validate file
        self._validate_file(file)

        # Generate unique filename
        file_extension = self._get_file_extension(file.filename)
        unique_filename = f"{uuid.uuid4().hex}_{teacher_id}.{file_extension}"

        # Create teacher-specific directory
        teacher_dir = self.UPLOAD_DIR / f"teacher_{teacher_id}"
        teacher_dir.mkdir(parents=True, exist_ok=True)

        file_path = teacher_dir / unique_filename

        # Save file
        file_size = 0
        try:
            content = await file.read()
            file_size = len(content)

            # Check file size
            if file_size > self.MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=413,
                    detail=f"File size exceeds maximum allowed size of {self.MAX_FILE_SIZE // (1024*1024)} MB"
                )

            with open(file_path, "wb") as buffer:
                buffer.write(content)

            return str(file_path), unique_filename, file_size

        except Exception as e:
            # Clean up file if something went wrong
            if os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")

    def _validate_file(self, file: UploadFile):
        """Validate file type and extension"""
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")

        # Check extension
        file_extension = self._get_file_extension(file.filename)
        if file_extension not in self.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"File type not allowed. Allowed types: {', '.join(self.ALLOWED_EXTENSIONS.keys())}"
            )

        # Check content type
        allowed_content_types = self.ALLOWED_EXTENSIONS[file_extension] + ["application/octet-stream"]
        if file.content_type not in allowed_content_types:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid content type '{file.content_type}' for {file_extension} file"
            )

    def _get_file_extension(self, filename: str) -> str:
        """Extract file extension"""
        if "." not in filename:
            raise HTTPException(status_code=400, detail="File has no extension")
        return filename.rsplit(".", 1)[-1].lower()

    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from PDF using pdfplumber (better formatting)
        Returns: (extracted_text, metadata)
        """
        try:
            import pdfplumber

            text_content = []
            metadata = {
                "pages": 0,
                "page_texts": {},
                "extraction_method": "pdfplumber"
            }

            with pdfplumber.open(file_path) as pdf:
                metadata["pages"] = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text() or ""
                    text_content.append(page_text)
                    metadata["page_texts"][page_num] = page_text

            full_text = "\n\n".join(text_content)
            return self._clean_text(full_text), metadata

        except ImportError:
            # Fallback to PyPDF2 if pdfplumber not installed
            print("pdfplumber not installed, trying PyPDF2...")
            return self._extract_text_from_pdf_pypdf2(file_path)
        except Exception as e:
            # Try PyPDF2 as fallback
            print(f"pdfplumber failed: {e}, trying PyPDF2...")
            return self._extract_text_from_pdf_pypdf2(file_path)

    def _extract_text_from_pdf_pypdf2(self, file_path: str) -> Tuple[str, Dict]:
        """Fallback PDF extraction using PyPDF2"""
        try:
            import PyPDF2

            text_content = []
            metadata = {
                "pages": 0,
                "page_texts": {},
                "extraction_method": "PyPDF2"
            }

            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    page_text = page.extract_text() or ""
                    text_content.append(page_text)
                    metadata["page_texts"][page_num] = page_text

            full_text = "\n\n".join(text_content)
            return self._clean_text(full_text), metadata

        except ImportError:
            raise Exception("Neither pdfplumber nor PyPDF2 is installed. Install with: pip install pdfplumber PyPDF2")
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from DOCX
        Returns: (extracted_text, metadata)
        """
        try:
            from docx import Document

            doc = Document(file_path)

            text_content = []
            metadata = {
                "paragraphs": 0,
                "tables": 0,
                "extraction_method": "python-docx"
            }

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
                    metadata["paragraphs"] += 1

            # Extract tables
            for table in doc.tables:
                metadata["tables"] += 1
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)

            full_text = "\n\n".join(text_content)
            return self._clean_text(full_text), metadata

        except ImportError:
            raise Exception("python-docx not installed. Install with: pip install python-docx")
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")

    def extract_text_from_txt(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from TXT
        Returns: (extracted_text, metadata)
        """
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                text = file.read()

            metadata = {
                "lines": len(text.split("\n")),
                "extraction_method": "plain_text",
                "encoding": "utf-8"
            }

            return self._clean_text(text), metadata

        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ["latin-1", "iso-8859-1", "cp1252"]:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        text = file.read()

                    metadata = {
                        "lines": len(text.split("\n")),
                        "extraction_method": "plain_text",
                        "encoding": encoding
                    }

                    return self._clean_text(text), metadata
                except:
                    continue

            raise Exception("Failed to decode text file with any common encoding")

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""

        # Remove excessive whitespace
        text = re.sub(r'[ \t]+', ' ', text)

        # Remove excessive newlines (keep max 2)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)

        # Remove special characters but keep punctuation
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)

        # Strip leading/trailing whitespace
        text = text.strip()

        return text

    async def process_file(
        self,
        file_path: str,
        file_type: str
    ) -> Tuple[str, Dict]:
        """
        Main processing function - extract text based on file type
        Returns: (extracted_text, metadata)
        """
        if file_type == "pdf":
            return self.extract_text_from_pdf(file_path)
        elif file_type == "docx":
            return self.extract_text_from_docx(file_path)
        elif file_type == "txt":
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def delete_file(self, file_path: str):
        """Delete file from disk"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"✅ Deleted file: {file_path}")
        except Exception as e:
            print(f"⚠️ Error deleting file {file_path}: {e}")
